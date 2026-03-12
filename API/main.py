from fastapi import FastAPI, UploadFile, File, HTTPException, Form, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pdf2docx import Converter
from pdf2image import convert_from_path
from docx import Document
from pypdf import PdfWriter, PdfReader
from typing import List
from concurrent.futures import ThreadPoolExecutor
import tempfile, os, shutil, zipfile, subprocess, pytesseract, camelot, uuid
from PIL import Image

app = FastAPI(title="Umbrella PDF Engine PRO")

# --- CONFIGURATION CORS ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Sur Render avec le Dockerfile nogui, 'soffice' est l'exécutable standard 
# qui pointe vers le binaire sans interface graphique.
LIBREOFFICE_BIN = "soffice"

# --- UTILITAIRES ---

def cleanup(temp_dir: str):
    """Supprime le dossier temporaire avec gestion d'erreurs (fichiers verrouillés)"""
    try:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
    except Exception as e:
        print(f"Erreur lors du nettoyage : {e}")

def process_pdf_to_word(pdf_path, docx_path):
    """Logique PDF -> Word avec fallback OCR et gestion de mémoire"""
    try:
        # 1. Tentative de conversion native (PDF textuel)
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()

        # 2. Détection de PDF scanné (Fallback OCR)
        # Si le fichier Word fait moins de 5KB, il est probablement vide ou contient juste des images sans texte
        if not os.path.exists(docx_path) or os.path.getsize(docx_path) < 5000:
            doc = Document()
            # On limite le DPI à 150 pour Render (512MB RAM) afin d'éviter le crash
            pages = convert_from_path(pdf_path, thread_count=1, dpi=150) 
            
            for i, page in enumerate(pages):
                # Utilisation de Tesseract pour extraire le texte
                text = pytesseract.image_to_string(page, lang='fra+eng')
                doc.add_paragraph(text)
                if i < len(pages) - 1:
                    doc.add_page_break()
                
                # Libération explicite de la mémoire de l'image pour Render
                page.close() 
            
            doc.save(docx_path)
            
        return docx_path
    except Exception as e:
        print(f"Erreur process_pdf_to_word: {e}")
        return None

@app.get("/")
def root():
    return {"status": "Umbrella Engine Online", "environment": "Production-Ready"}

# --- CATEGORIE : CONVERT ---

@app.post("/convert/ocr")
async def ocr_pdf(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    try:
        zip_path = os.path.join(temp_dir, "umbrella_ocr.zip")
        with zipfile.ZipFile(zip_path, "w") as z:
            for file in files:
                safe_name = f"{uuid.uuid4()}_{os.path.basename(file.filename)}"
                p = os.path.join(temp_dir, safe_name)
                with open(p, "wb") as f:
                    shutil.copyfileobj(file.file, f)
                await file.close()

                # DPI=150 pour préserver la RAM sur Render
                pages = convert_from_path(p, thread_count=1, dpi=150)
                doc = Document()
                for page in pages:
                    text = pytesseract.image_to_string(page, lang='fra+eng')
                    doc.add_paragraph(text)
                    page.close() # CRITIQUE : Libère la mémoire !
                
                out_docx = p.replace(".pdf", "_ocr.docx")
                doc.save(out_docx)
                z.write(out_docx, os.path.basename(out_docx))
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(zip_path, filename="umbrella_ocr_results.zip", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/pdf-to-word")
async def pdf_to_word(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    try:
        zip_path = os.path.join(temp_dir, "umbrella_word.zip")
        tasks = []
        for file in files:
            safe_name = f"{uuid.uuid4()}_{os.path.basename(file.filename)}"
            p = os.path.join(temp_dir, safe_name)
            d = p.replace(".pdf", ".docx")
            with open(p, "wb") as f:
                shutil.copyfileobj(file.file, f)
            await file.close()
            tasks.append((p, d))
        
        with ThreadPoolExecutor() as ex:
            list(ex.map(lambda t: process_pdf_to_word(*t), tasks))
        
        with zipfile.ZipFile(zip_path, "w") as z:
            for f in os.listdir(temp_dir):
                if f.endswith(".docx"):
                    z.write(os.path.join(temp_dir, f), f)
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(zip_path, filename="umbrella_word.zip", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/office-to-pdf")
async def office_to_pdf(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    # On crée un dossier spécifique pour le profil LibreOffice pour éviter les conflits
    user_profile = os.path.join(temp_dir, "profile")
    try:
        for file in files:
            safe_name = f"{uuid.uuid4()}_{os.path.basename(file.filename)}"
            in_p = os.path.join(temp_dir, safe_name)
            with open(in_p, "wb") as f: 
                shutil.copyfileobj(file.file, f)
            await file.close()
            
            # Commande blindée pour serveur headless
            subprocess.run([
                LIBREOFFICE_BIN, 
                "--headless", 
                "--nodefault",
                "--nofirststartwizard",
                f"-env:UserInstallation=file://{user_profile}",
                "--convert-to", "pdf", 
                in_p, 
                "--outdir", temp_dir
            ], check=True, timeout=60)

        zip_path = os.path.join(temp_dir, "umbrella_office.zip")
        with zipfile.ZipFile(zip_path, "w") as z:
            for f in os.listdir(temp_dir):
                if f.endswith(".pdf"): 
                    z.write(os.path.join(temp_dir, f), f)
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(zip_path, filename="umbrella_office_converted.zip", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))
@app.post("/convert/pdf-to-excel")
async def pdf_to_excel(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    zip_path = os.path.join(temp_dir, "umbrella_excel.zip")
    try:
        with zipfile.ZipFile(zip_path, "w") as z:
            for file in files:
                safe_name = f"{uuid.uuid4()}_{os.path.basename(file.filename)}"
                p = os.path.join(temp_dir, safe_name)
                with open(p, "wb") as f: 
                    shutil.copyfileobj(file.file, f)
                await file.close()
                out_xlsx = p.replace(".pdf", ".xlsx")
                tables = camelot.read_pdf(p, pages="all", flavor='stream') # 'stream' est souvent plus stable sur serveur
                tables.export(out_xlsx, f="excel")
                if os.path.exists(out_xlsx): 
                    z.write(out_xlsx, os.path.basename(out_xlsx))
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(zip_path, filename="umbrella_excel.zip", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/pdf-to-jpg")
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pdf2docx import Converter
from pdf2image import convert_from_path
from docx import Document
from pypdf import PdfWriter, PdfReader
from typing import List
from concurrent.futures import ThreadPoolExecutor
import tempfile, os, shutil, zipfile, subprocess, pytesseract, uuid
import camelot
from PIL import Image

app = FastAPI(title="Umbrella PDF Engine PRO")

# --- CONFIGURATION ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

LIBREOFFICE_BIN = "soffice"

# --- UTILITAIRES ---

def cleanup(temp_dir: str):
    """Nettoyage sécurisé des fichiers temporaires après réponse"""
    try:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
    except Exception as e:
        print(f"Erreur nettoyage : {e}")

def process_pdf_to_word(pdf_path, docx_path):
    """Conversion PDF vers Word avec Fallback OCR intelligent"""
    try:
        # 1. Conversion native
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()

        # 2. Fallback OCR si le fichier est vide/scanné (seuil 5KB)
        if not os.path.exists(docx_path) or os.path.getsize(docx_path) < 5000:
            doc = Document()
            # DPI=100 pour préserver la RAM sur Render
            pages = convert_from_path(pdf_path, thread_count=1, dpi=100) 
            for page in pages:
                text = pytesseract.image_to_string(page, lang='fra+eng')
                doc.add_paragraph(text)
                page.close()
            doc.save(docx_path)
        return docx_path
    except Exception as e:
        print(f"Erreur PDF2Word: {e}")
        return None

@app.get("/")
def root():
    return {"status": "Umbrella Engine Online", "environment": "Render-Ready"}

# --- CATEGORIE : CONVERSION ---

@app.post("/convert/office-to-pdf")
async def office_to_pdf(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    user_profile = os.path.join(temp_dir, "profile")
    try:
        for file in files:
            safe_name = f"{uuid.uuid4()}_{file.filename.replace(' ', '_')}"
            in_p = os.path.join(temp_dir, safe_name)
            with open(in_p, "wb") as f: 
                shutil.copyfileobj(file.file, f)
            
            # Commande durcie pour LibreOffice Headless
            subprocess.run([
                LIBREOFFICE_BIN, "--headless", "--nologo", "--nodefault",
                f"-env:UserInstallation=file://{os.path.abspath(user_profile)}",
                "--convert-to", "pdf", in_p, "--outdir", temp_dir
            ], check=True, timeout=120)

        zip_path = os.path.join(temp_dir, "office_converted.zip")
        with zipfile.ZipFile(zip_path, "w") as z:
            for f in os.listdir(temp_dir):
                if f.lower().endswith(".pdf"): 
                    z.write(os.path.join(temp_dir, f), f)
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(zip_path, filename="office_converted.zip", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/pdf-to-word")
async def pdf_to_word(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    try:
        tasks = []
        for file in files:
            p = os.path.join(temp_dir, f"{uuid.uuid4()}_{file.filename}")
            d = p.rsplit(".", 1)[0] + ".docx"
            with open(p, "wb") as f:
                shutil.copyfileobj(file.file, f)
            tasks.append((p, d))
        
        # Max 2 workers pour ne pas exploser la RAM Render
        with ThreadPoolExecutor(max_workers=2) as ex:
            list(ex.map(lambda t: process_pdf_to_word(*t), tasks))
        
        zip_path = os.path.join(temp_dir, "pdf_to_word.zip")
        with zipfile.ZipFile(zip_path, "w") as z:
            for f in os.listdir(temp_dir):
                if f.endswith(".docx"): z.write(os.path.join(temp_dir, f), f)
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(zip_path, filename="pdf_to_word.zip", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/pdf-to-jpg")
async def pdf_to_jpg(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    temp_dir = tempfile.mkdtemp()
    try:
        p = os.path.join(temp_dir, f"{uuid.uuid4()}_{file.filename}")
        with open(p, "wb") as f: shutil.copyfileobj(file.file, f)
        
        # paths_only=True : Écrit sur disque au lieu de charger en RAM
        image_paths = convert_from_path(p, dpi=150, output_folder=temp_dir, fmt="jpg", paths_only=True)
        
        zip_path = os.path.join(temp_dir, "images.zip")
        with zipfile.ZipFile(zip_path, "w") as z:
            for i, path in enumerate(image_paths):
                z.write(path, f"page_{i+1}.jpg")
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(zip_path, filename="pdf_images.zip", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/convert/images-to-pdf")
async def images_to_pdf(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    output_pdf = os.path.join(temp_dir, "merged_images.pdf")
    try:
        img_objs = []
        for file in files:
            path = os.path.join(temp_dir, f"{uuid.uuid4()}_{file.filename}")
            with open(path, "wb") as f: shutil.copyfileobj(file.file, f)
            img_objs.append(Image.open(path).convert("RGB"))
        
        if img_objs:
            img_objs[0].save(output_pdf, save_all=True, append_images=img_objs[1:])
            for img in img_objs: img.close() # Libère la RAM
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(output_pdf, filename="images_to_pdf.pdf", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

# --- CATEGORIE : ORGANISATION & SECURITE ---

@app.post("/organize/merge")
async def merge_pdfs(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    temp_dir = tempfile.mkdtemp()
    merger = PdfWriter()
    try:
        for file in files:
            p = os.path.join(temp_dir, f"{uuid.uuid4()}_{file.filename}")
            with open(p, "wb") as f: shutil.copyfileobj(file.file, f)
            merger.append(p)
        
        out = os.path.join(temp_dir, "merged.pdf")
        merger.write(out)
        merger.close()
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(out, filename="merged.pdf", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/security/protect")
async def protect_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...), password: str = Form(...)):
    temp_dir = tempfile.mkdtemp()
    try:
        p = os.path.join(temp_dir, file.filename)
        with open(p, "wb") as f: shutil.copyfileobj(file.file, f)
        
        reader = PdfReader(p)
        writer = PdfWriter()
        for page in reader.pages: writer.add_page(page)
        writer.encrypt(password)
        
        out = os.path.join(temp_dir, "protected.pdf")
        with open(out, "wb") as f: writer.write(f)
        
        background_tasks.add_task(cleanup, temp_dir)
        return FileResponse(out, filename="protected.pdf", background=background_tasks)
    except Exception as e:
        cleanup(temp_dir)
        raise HTTPException(status_code=500, detail=str(e))
